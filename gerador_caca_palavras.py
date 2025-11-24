#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Gerador de Caça-Palavras em PDF, JPEG e DOCX
By Ramon Las-cazas

BIBLIOTECAS NECESSÁRIAS:
pip install reportlab Pillow python-docx

Se alguma biblioteca estiver faltando, instale usando o comando acima.
"""

import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
import random
import string
import os
import unicodedata

# Importações obrigatórias
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors
except ImportError:
    messagebox.showerror("Erro de Importação", 
                        "Biblioteca 'reportlab' não encontrada!\n\n"
                        "Instale usando: pip install reportlab")
    exit()

# Importações opcionais para JPEG
try:
    from PIL import Image, ImageDraw, ImageFont
    PILLOW_DISPONIVEL = True
except ImportError:
    PILLOW_DISPONIVEL = False

# Importações opcionais para DOCX
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False

class GeradorCacaPalavras:
    def __init__(self):
        self.grade = []
        self.palavras_posicoes = []
        self.tamanho = 0
        
    def remover_acentos(self, texto):
        """Remove acentos de uma string"""
        nfkd = unicodedata.normalize('NFKD', texto)
        return ''.join([c for c in nfkd if not unicodedata.combining(c)])
    
    def criar_grade_vazia(self, tamanho):
        """Cria uma grade vazia preenchida com None"""
        self.tamanho = tamanho
        self.grade = [[None for _ in range(tamanho)] for _ in range(tamanho)]
        self.palavras_posicoes = []
    
    def pode_colocar_palavra(self, palavra, linha, coluna, direcao):
        """Verifica se é possível colocar a palavra na posição e direção especificadas"""
        palavra_sem_acento = self.remover_acentos(palavra.upper()).replace(" ", "")
        tamanho_palavra = len(palavra_sem_acento)
        
        # Direções: 
        # 0=horizontal direita, 1=horizontal esquerda, 
        # 2=vertical baixo, 3=vertical cima
        # 4=diagonal baixo-direita, 5=diagonal baixo-esquerda
        # 6=diagonal cima-direita, 7=diagonal cima-esquerda
        
        if direcao == 0:  # Horizontal direita
            if coluna + tamanho_palavra > self.tamanho:
                return False
            for i in range(tamanho_palavra):
                if self.grade[linha][coluna + i] is not None and self.grade[linha][coluna + i] != palavra_sem_acento[i]:
                    return False
        
        elif direcao == 1:  # Horizontal esquerda
            if coluna - tamanho_palavra + 1 < 0:
                return False
            for i in range(tamanho_palavra):
                if self.grade[linha][coluna - i] is not None and self.grade[linha][coluna - i] != palavra_sem_acento[i]:
                    return False
        
        elif direcao == 2:  # Vertical para baixo
            if linha + tamanho_palavra > self.tamanho:
                return False
            for i in range(tamanho_palavra):
                if self.grade[linha + i][coluna] is not None and self.grade[linha + i][coluna] != palavra_sem_acento[i]:
                    return False
        
        elif direcao == 3:  # Vertical para cima
            if linha - tamanho_palavra + 1 < 0:
                return False
            for i in range(tamanho_palavra):
                if self.grade[linha - i][coluna] is not None and self.grade[linha - i][coluna] != palavra_sem_acento[i]:
                    return False
        
        elif direcao == 4:  # Diagonal baixo-direita
            if linha + tamanho_palavra > self.tamanho or coluna + tamanho_palavra > self.tamanho:
                return False
            for i in range(tamanho_palavra):
                if self.grade[linha + i][coluna + i] is not None and self.grade[linha + i][coluna + i] != palavra_sem_acento[i]:
                    return False
        
        elif direcao == 5:  # Diagonal baixo-esquerda
            if linha + tamanho_palavra > self.tamanho or coluna - tamanho_palavra + 1 < 0:
                return False
            for i in range(tamanho_palavra):
                if self.grade[linha + i][coluna - i] is not None and self.grade[linha + i][coluna - i] != palavra_sem_acento[i]:
                    return False
        
        elif direcao == 6:  # Diagonal cima-direita
            if linha - tamanho_palavra + 1 < 0 or coluna + tamanho_palavra > self.tamanho:
                return False
            for i in range(tamanho_palavra):
                if self.grade[linha - i][coluna + i] is not None and self.grade[linha - i][coluna + i] != palavra_sem_acento[i]:
                    return False
        
        elif direcao == 7:  # Diagonal cima-esquerda
            if linha - tamanho_palavra + 1 < 0 or coluna - tamanho_palavra + 1 < 0:
                return False
            for i in range(tamanho_palavra):
                if self.grade[linha - i][coluna - i] is not None and self.grade[linha - i][coluna - i] != palavra_sem_acento[i]:
                    return False
        
        return True
    
    def colocar_palavra(self, palavra, linha, coluna, direcao):
        """Coloca a palavra na grade"""
        palavra_sem_acento = self.remover_acentos(palavra.upper()).replace(" ", "")
        tamanho_palavra = len(palavra_sem_acento)
        posicoes = []
        
        if direcao == 0:  # Horizontal direita
            for i in range(tamanho_palavra):
                self.grade[linha][coluna + i] = palavra_sem_acento[i]
                posicoes.append((linha, coluna + i))
        
        elif direcao == 1:  # Horizontal esquerda
            for i in range(tamanho_palavra):
                self.grade[linha][coluna - i] = palavra_sem_acento[i]
                posicoes.append((linha, coluna - i))
        
        elif direcao == 2:  # Vertical para baixo
            for i in range(tamanho_palavra):
                self.grade[linha + i][coluna] = palavra_sem_acento[i]
                posicoes.append((linha + i, coluna))
        
        elif direcao == 3:  # Vertical para cima
            for i in range(tamanho_palavra):
                self.grade[linha - i][coluna] = palavra_sem_acento[i]
                posicoes.append((linha - i, coluna))
        
        elif direcao == 4:  # Diagonal baixo-direita
            for i in range(tamanho_palavra):
                self.grade[linha + i][coluna + i] = palavra_sem_acento[i]
                posicoes.append((linha + i, coluna + i))
        
        elif direcao == 5:  # Diagonal baixo-esquerda
            for i in range(tamanho_palavra):
                self.grade[linha + i][coluna - i] = palavra_sem_acento[i]
                posicoes.append((linha + i, coluna - i))
        
        elif direcao == 6:  # Diagonal cima-direita
            for i in range(tamanho_palavra):
                self.grade[linha - i][coluna + i] = palavra_sem_acento[i]
                posicoes.append((linha - i, coluna + i))
        
        elif direcao == 7:  # Diagonal cima-esquerda
            for i in range(tamanho_palavra):
                self.grade[linha - i][coluna - i] = palavra_sem_acento[i]
                posicoes.append((linha - i, coluna - i))
        
        self.palavras_posicoes.append({
            'palavra': palavra,
            'posicoes': posicoes
        })
    
    def inserir_palavras(self, palavras, usar_diagonais=False, usar_contrarias=True):
        """Tenta inserir todas as palavras na grade"""
        palavras_nao_inseridas = []
        
        # Definir direções disponíveis
        if usar_diagonais:
            if usar_contrarias:
                # Todas as direções (0-7)
                direcoes_disponiveis = list(range(8))
            else:
                # Apenas direções "para frente": horizontal direita, vertical baixo, diagonais para baixo
                direcoes_disponiveis = [0, 2, 4, 5]
        else:
            if usar_contrarias:
                # Apenas horizontal e vertical (todas)
                direcoes_disponiveis = [0, 1, 2, 3]
            else:
                # Apenas horizontal direita e vertical para baixo
                direcoes_disponiveis = [0, 2]
        
        for palavra in palavras:
            palavra_limpa = palavra.strip()
            if not palavra_limpa:
                continue
                
            inserida = False
            tentativas = 0
            max_tentativas = 100
            
            while not inserida and tentativas < max_tentativas:
                linha = random.randint(0, self.tamanho - 1)
                coluna = random.randint(0, self.tamanho - 1)
                direcao = random.choice(direcoes_disponiveis)
                
                if self.pode_colocar_palavra(palavra_limpa, linha, coluna, direcao):
                    self.colocar_palavra(palavra_limpa, linha, coluna, direcao)
                    inserida = True
                
                tentativas += 1
            
            if not inserida:
                palavras_nao_inseridas.append(palavra_limpa)
        
        return palavras_nao_inseridas
    
    def preencher_espacos_vazios(self):
        """Preenche os espaços vazios com letras aleatórias"""
        letras = string.ascii_uppercase
        for i in range(self.tamanho):
            for j in range(self.tamanho):
                if self.grade[i][j] is None:
                    self.grade[i][j] = random.choice(letras)
    
    def gerar_pdf(self, nome_arquivo, palavras_originais):
        """Gera o PDF com o caça-palavras e o gabarito"""
        c = canvas.Canvas(nome_arquivo, pagesize=A4)
        largura, altura = A4
        
        # Calcular tamanho das células
        margem = 50
        espaco_disponivel = min(largura - 2 * margem, altura - 200)
        tamanho_celula = espaco_disponivel / self.tamanho
        
        # PÁGINA 1: CAÇA-PALAVRAS
        c.setFont("Helvetica-Bold", 20)
        c.drawCentredString(largura / 2, altura - 30, "CAÇA-PALAVRAS")
        
        # Desenhar a grade
        inicio_x = (largura - (tamanho_celula * self.tamanho)) / 2
        inicio_y = altura - 100
        
        c.setFont("Helvetica", int(tamanho_celula * 0.6))
        
        for i in range(self.tamanho):
            for j in range(self.tamanho):
                x = inicio_x + j * tamanho_celula
                y = inicio_y - i * tamanho_celula
                
                # Desenhar borda da célula
                c.rect(x, y - tamanho_celula, tamanho_celula, tamanho_celula)
                
                # Desenhar letra centralizada
                letra = self.grade[i][j]
                texto_x = x + tamanho_celula / 2
                texto_y = y - tamanho_celula / 2 - int(tamanho_celula * 0.2)
                c.drawCentredString(texto_x, texto_y, letra)
        
        # Lista de palavras
        c.setFont("Helvetica-Bold", 14)
        y_palavras = inicio_y - (self.tamanho * tamanho_celula) - 30
        c.drawString(margem, y_palavras, "PALAVRAS:")
        
        c.setFont("Helvetica", 11)
        y_atual = y_palavras - 20
        x_coluna1 = margem
        x_coluna2 = largura / 2
        coluna_atual = 0
        
        for idx, palavra in enumerate(palavras_originais):
            if coluna_atual == 0:
                c.drawString(x_coluna1, y_atual, f"• {palavra}")
            else:
                c.drawString(x_coluna2, y_atual, f"• {palavra}")
            
            coluna_atual = 1 - coluna_atual
            if coluna_atual == 0:
                y_atual -= 15
        
        # PÁGINA 2: GABARITO
        c.showPage()
        
        c.setFont("Helvetica-Bold", 20)
        c.drawCentredString(largura / 2, altura - 30, "GABARITO")
        
        # Desenhar a grade novamente
        c.setFont("Helvetica", int(tamanho_celula * 0.6))
        
        for i in range(self.tamanho):
            for j in range(self.tamanho):
                x = inicio_x + j * tamanho_celula
                y = inicio_y - i * tamanho_celula
                
                # Desenhar borda da célula
                c.rect(x, y - tamanho_celula, tamanho_celula, tamanho_celula)
                
                # Desenhar letra
                letra = self.grade[i][j]
                texto_x = x + tamanho_celula / 2
                texto_y = y - tamanho_celula / 2 - int(tamanho_celula * 0.2)
                c.drawCentredString(texto_x, texto_y, letra)
        
        # Destacar palavras com cores diferentes
        cores_disponiveis = [
            colors.red, colors.blue, colors.green, colors.orange,
            colors.purple, colors.brown, colors.pink, colors.cyan,
            colors.magenta, colors.yellow, colors.lightblue, colors.lightgreen
        ]
        
        for idx, palavra_info in enumerate(self.palavras_posicoes):
            cor = cores_disponiveis[idx % len(cores_disponiveis)]
            c.setStrokeColor(cor)
            c.setLineWidth(3)
            
            posicoes = palavra_info['posicoes']
            if len(posicoes) > 0:
                # Calcular retângulo envolvente
                linhas = [pos[0] for pos in posicoes]
                colunas = [pos[1] for pos in posicoes]
                
                min_linha = min(linhas)
                max_linha = max(linhas)
                min_coluna = min(colunas)
                max_coluna = max(colunas)
                
                x1 = inicio_x + min_coluna * tamanho_celula
                y1 = inicio_y - min_linha * tamanho_celula
                largura_ret = (max_coluna - min_coluna + 1) * tamanho_celula
                altura_ret = (max_linha - min_linha + 1) * tamanho_celula
                
                c.rect(x1, y1 - altura_ret, largura_ret, altura_ret, stroke=1, fill=0)
        
        c.save()
        return True
    
    def gerar_jpeg(self, nome_arquivo, palavras_originais, incluir_gabarito=True):
        """Gera imagem JPEG com o caça-palavras e opcionalmente o gabarito"""
        if not PILLOW_DISPONIVEL:
            raise ImportError("Biblioteca PIL/Pillow não está instalada. Use: pip install Pillow")
        
        # Dimensões da imagem
        largura_img = 2480  # A4 em 300 DPI
        altura_img = 3508
        
        # Criar imagem para caça-palavras
        img = Image.new('RGB', (largura_img, altura_img), 'white')
        draw = ImageDraw.Draw(img)
        
        # Tentar carregar fonte, usar padrão se não encontrar
        try:
            fonte_titulo = ImageFont.truetype("arial.ttf", 80)
            fonte_celula = ImageFont.truetype("arial.ttf", int(1500 / self.tamanho))
            fonte_palavra = ImageFont.truetype("arial.ttf", 40)
        except:
            fonte_titulo = ImageFont.load_default()
            fonte_celula = ImageFont.load_default()
            fonte_palavra = ImageFont.load_default()
        
        # Título
        titulo = "CAÇA-PALAVRAS"
        bbox = draw.textbbox((0, 0), titulo, font=fonte_titulo)
        titulo_largura = bbox[2] - bbox[0]
        draw.text((largura_img/2 - titulo_largura/2, 100), titulo, fill='black', font=fonte_titulo)
        
        # Calcular dimensões da grade
        margem = 200
        espaco_disponivel = min(largura_img - 2 * margem, altura_img - 800)
        tamanho_celula = espaco_disponivel / self.tamanho
        
        inicio_x = (largura_img - (tamanho_celula * self.tamanho)) / 2
        inicio_y = 300
        
        # Desenhar grade
        for i in range(self.tamanho):
            for j in range(self.tamanho):
                x = inicio_x + j * tamanho_celula
                y = inicio_y + i * tamanho_celula
                
                # Desenhar borda
                draw.rectangle([x, y, x + tamanho_celula, y + tamanho_celula], outline='black', width=2)
                
                # Desenhar letra
                letra = self.grade[i][j]
                bbox = draw.textbbox((0, 0), letra, font=fonte_celula)
                letra_largura = bbox[2] - bbox[0]
                letra_altura = bbox[3] - bbox[1]
                texto_x = x + tamanho_celula/2 - letra_largura/2
                texto_y = y + tamanho_celula/2 - letra_altura/2
                draw.text((texto_x, texto_y), letra, fill='black', font=fonte_celula)
        
        # Lista de palavras
        y_palavras = inicio_y + (self.tamanho * tamanho_celula) + 80
        draw.text((margem, y_palavras), "PALAVRAS:", fill='black', font=fonte_palavra)
        
        y_atual = y_palavras + 60
        x_coluna1 = margem
        x_coluna2 = largura_img / 2
        coluna_atual = 0
        
        for palavra in palavras_originais:
            if coluna_atual == 0:
                draw.text((x_coluna1, y_atual), f"• {palavra}", fill='black', font=fonte_palavra)
            else:
                draw.text((x_coluna2, y_atual), f"• {palavra}", fill='black', font=fonte_palavra)
            
            coluna_atual = 1 - coluna_atual
            if coluna_atual == 0:
                y_atual += 50
        
        # Salvar primeira imagem
        if incluir_gabarito:
            base_nome = nome_arquivo.rsplit('.', 1)[0]
            img.save(f"{base_nome}_caca.jpeg", 'JPEG', quality=95)
            
            # Criar imagem do gabarito
            img_gab = Image.new('RGB', (largura_img, altura_img), 'white')
            draw_gab = ImageDraw.Draw(img_gab)
            
            # Título gabarito
            titulo_gab = "GABARITO"
            bbox = draw_gab.textbbox((0, 0), titulo_gab, font=fonte_titulo)
            titulo_largura = bbox[2] - bbox[0]
            draw_gab.text((largura_img/2 - titulo_largura/2, 100), titulo_gab, fill='black', font=fonte_titulo)
            
            # Desenhar grade
            for i in range(self.tamanho):
                for j in range(self.tamanho):
                    x = inicio_x + j * tamanho_celula
                    y = inicio_y + i * tamanho_celula
                    
                    draw_gab.rectangle([x, y, x + tamanho_celula, y + tamanho_celula], outline='black', width=2)
                    
                    letra = self.grade[i][j]
                    bbox = draw_gab.textbbox((0, 0), letra, font=fonte_celula)
                    letra_largura = bbox[2] - bbox[0]
                    letra_altura = bbox[3] - bbox[1]
                    texto_x = x + tamanho_celula/2 - letra_largura/2
                    texto_y = y + tamanho_celula/2 - letra_altura/2
                    draw_gab.text((texto_x, texto_y), letra, fill='black', font=fonte_celula)
            
            # Destacar palavras
            cores_rgb = [
                (255, 0, 0), (0, 0, 255), (0, 128, 0), (255, 165, 0),
                (128, 0, 128), (165, 42, 42), (255, 192, 203), (0, 255, 255)
            ]
            
            for idx, palavra_info in enumerate(self.palavras_posicoes):
                cor = cores_rgb[idx % len(cores_rgb)]
                posicoes = palavra_info['posicoes']
                
                if len(posicoes) > 0:
                    linhas = [pos[0] for pos in posicoes]
                    colunas = [pos[1] for pos in posicoes]
                    
                    min_linha = min(linhas)
                    max_linha = max(linhas)
                    min_coluna = min(colunas)
                    max_coluna = max(colunas)
                    
                    x1 = inicio_x + min_coluna * tamanho_celula
                    y1 = inicio_y + min_linha * tamanho_celula
                    x2 = inicio_x + (max_coluna + 1) * tamanho_celula
                    y2 = inicio_y + (max_linha + 1) * tamanho_celula
                    
                    draw_gab.rectangle([x1, y1, x2, y2], outline=cor, width=5)
            
            img_gab.save(f"{base_nome}_gabarito.jpeg", 'JPEG', quality=95)
        else:
            img.save(nome_arquivo, 'JPEG', quality=95)
        
        return True
    
    def gerar_docx(self, nome_arquivo, palavras_originais):
        """Gera documento DOCX com o caça-palavras e o gabarito em tabelas"""
        if not DOCX_DISPONIVEL:
            raise ImportError("Biblioteca python-docx não está instalada. Use: pip install python-docx")
        
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # PÁGINA 1: CAÇA-PALAVRAS
        titulo = doc.add_heading('CAÇA-PALAVRAS', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Criar tabela para a grade
        tabela = doc.add_table(rows=self.tamanho, cols=self.tamanho)
        tabela.style = 'Table Grid'
        tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Calcular tamanho ideal das células
        largura_celula = 6.5 / self.tamanho  # Total de 6.5 polegadas
        
        # Configurar células
        for i in range(self.tamanho):
            row = tabela.rows[i]
            row.height = Inches(largura_celula)  # Células quadradas
            
            for j in range(self.tamanho):
                celula = row.cells[j]
                celula.width = Inches(largura_celula)
                celula.text = self.grade[i][j]
                
                # Centralizar texto vertical e horizontalmente
                celula.vertical_alignment = 1  # 1 = CENTER
                paragrafo = celula.paragraphs[0]
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Formatar fonte
                run = paragrafo.runs[0]
                tamanho_fonte = max(8, min(16, int(200 / self.tamanho)))
                run.font.size = Pt(tamanho_fonte)
                run.font.bold = True
                run.font.name = 'Arial'
                
                # Ajustar espaçamento do parágrafo
                paragrafo.paragraph_format.space_before = Pt(0)
                paragrafo.paragraph_format.space_after = Pt(0)
        
        # Lista de palavras
        doc.add_paragraph()
        palavras_titulo = doc.add_heading('PALAVRAS:', level=2)
        
        # Criar tabela para as palavras (2 colunas)
        num_linhas = (len(palavras_originais) + 1) // 2
        tabela_palavras = doc.add_table(rows=num_linhas, cols=2)
        tabela_palavras.style = 'Light List'
        
        for i, palavra in enumerate(palavras_originais):
            linha = i // 2
            coluna = i % 2
            celula = tabela_palavras.rows[linha].cells[coluna]
            celula.text = f"• {palavra}"
            paragrafo = celula.paragraphs[0]
            run = paragrafo.runs[0]
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        
        # PÁGINA 2: GABARITO
        doc.add_page_break()
        
        titulo_gab = doc.add_heading('GABARITO', 0)
        titulo_gab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Criar tabela para o gabarito
        tabela_gab = doc.add_table(rows=self.tamanho, cols=self.tamanho)
        tabela_gab.style = 'Table Grid'
        tabela_gab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Marcar células que fazem parte das palavras
        celulas_palavras = set()
        for palavra_info in self.palavras_posicoes:
            for pos in palavra_info['posicoes']:
                celulas_palavras.add(pos)
        
        # Configurar células do gabarito
        for i in range(self.tamanho):
            row = tabela_gab.rows[i]
            row.height = Inches(largura_celula)
            
            for j in range(self.tamanho):
                celula = row.cells[j]
                celula.width = Inches(largura_celula)
                celula.text = self.grade[i][j]
                
                # Centralizar texto
                celula.vertical_alignment = 1
                paragrafo = celula.paragraphs[0]
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Formatar fonte
                run = paragrafo.runs[0]
                run.font.size = Pt(tamanho_fonte)
                run.font.bold = True
                run.font.name = 'Arial'
                
                # Destacar células com palavras em vermelho
                if (i, j) in celulas_palavras:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    # Destacar fundo da célula
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'FFFF00')  # Amarelo
                    celula._element.get_or_add_tcPr().append(shading_elm)
                
                # Ajustar espaçamento
                paragrafo.paragraph_format.space_before = Pt(0)
                paragrafo.paragraph_format.space_after = Pt(0)
        
        doc.save(nome_arquivo)
        return True


class InterfaceApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Gerador de Caça-Palavras")
        self.root.geometry("600x750")
        self.root.resizable(False, False)
        
        # Variável para armazenar o diretório selecionado
        self.diretorio_destino = os.path.expanduser("~")  # Diretório padrão: pasta do usuário
        
        # Variável para tema escuro
        self.tema_escuro = tk.BooleanVar(value=False)
        
        # Cores dos temas
        self.cores_claro = {
            'bg': '#f0f0f0',
            'fg': '#000000',
            'entry_bg': 'white',
            'entry_fg': '#000000',
            'button_bg': '#4CAF50',
            'button_fg': 'white',
            'frame_bg': '#f0f0f0',
            'label_info': '#FF0000',
            'label_link': 'blue',
            'border': '#d0d0d0'
        }
        
        self.cores_escuro = {
            'bg': '#2b2b2b',
            'fg': '#e0e0e0',
            'entry_bg': '#3c3c3c',
            'entry_fg': '#e0e0e0',
            'button_bg': '#45a049',
            'button_fg': 'white',
            'frame_bg': '#2b2b2b',
            'label_info': "#FF0000",
            'label_link': '#6ba3ff',
            'border': '#3c3c3c'
        }
        
        # Frame superior com título e tema
        frame_topo = tk.Frame(root)
        frame_topo.pack(fill=tk.X, padx=20, pady=5)
        
        # Título (centralizado)
        frame_titulo = tk.Frame(frame_topo)
        frame_titulo.pack(expand=True)
        
        self.titulo = tk.Label(frame_titulo, text="Gerador de Caça-Palavras em PDF", 
                         font=("Arial", 16, "bold"), pady=5)
        self.titulo.pack()
        
        self.subtitulo = tk.Label(frame_titulo, text="By Ramon Las-cazas", 
                         font=("Arial", 10), pady=5)
        self.subtitulo.pack()
        
        # Botão de tema no canto superior direito
        self.btn_tema = tk.Checkbutton(root, text="🌙 Tema Escuro", 
                                       variable=self.tema_escuro,
                                       command=self.alternar_tema,
                                       font=("Arial", 9),
                                       relief=tk.FLAT,
                                       highlightthickness=0,
                                       bd=0)
        self.btn_tema.place(x=480, y=10)
        
        # Armazenar referência ao frame_topo para aplicar tema
        self.frame_topo = frame_topo
        self.frame_titulo = frame_titulo

        # Frame para tamanho da grade
        self.frame_tamanho = tk.Frame(root, pady=5)
        self.frame_tamanho.pack(fill=tk.X, padx=20)
        
        self.label_tamanho = tk.Label(self.frame_tamanho, text="Tamanho da Grade:", 
                font=("Arial", 10))
        self.label_tamanho.pack(side=tk.LEFT)
        
        self.entry_tamanho = tk.Entry(self.frame_tamanho, width=10, font=("Arial", 10))
        self.entry_tamanho.pack(side=tk.LEFT, padx=10)
        self.entry_tamanho.insert(0, "18x18")
        
        self.label_exemplo_tamanho = tk.Label(self.frame_tamanho, text="(ex: 18x18, 15x15)", 
                font=("Arial", 9))
        self.label_exemplo_tamanho.pack(side=tk.LEFT)
        
        # Frame para opções
        self.frame_opcoes = tk.LabelFrame(root, text="Opções de Direção", 
                                      font=("Arial", 10, "bold"), pady=10, padx=10,
                                      relief=tk.GROOVE, bd=2)
        self.frame_opcoes.pack(fill=tk.X, padx=20, pady=10)
        
        # Opção 1: Incluir diagonais
        self.var_diagonais = tk.BooleanVar(value=False)
        self.check_diagonais = tk.Checkbutton(self.frame_opcoes, 
                                         text="Incluir palavras na diagonal",
                                         variable=self.var_diagonais,
                                         font=("Arial", 10))
        self.check_diagonais.pack(anchor=tk.W, pady=5)
        
        # Opção 2: Palavras ao contrário
        self.var_contrarias = tk.BooleanVar(value=True)
        self.check_contrarias = tk.Checkbutton(self.frame_opcoes, 
                                          text="Permitir palavras ao contrário",
                                          variable=self.var_contrarias,
                                          font=("Arial", 10))
        self.check_contrarias.pack(anchor=tk.W, pady=5)
        
        # Label explicativo
        self.label_explicativo = tk.Label(self.frame_opcoes, 
                text="• Sem diagonal: apenas horizontal e vertical\n• Ao contrário: palavras podem aparecer invertidas",
                font=("Arial", 8), justify=tk.LEFT)
        self.label_explicativo.pack(anchor=tk.W, padx=20)
        
        # Frame para palavras
        self.frame_palavras = tk.Frame(root, pady=5)
        self.frame_palavras.pack(fill=tk.BOTH, expand=True, padx=20)
        
        self.label_palavras = tk.Label(self.frame_palavras, text="Palavras (uma por linha):", 
                font=("Arial", 10))
        self.label_palavras.pack(anchor=tk.W)
        
        self.text_palavras = scrolledtext.ScrolledText(self.frame_palavras, 
                                                       height=10, 
                                                       font=("Arial", 10),
                                                       wrap=tk.WORD)
        self.text_palavras.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Frame para nome do arquivo
        self.frame_arquivo = tk.Frame(root, pady=5)
        self.frame_arquivo.pack(fill=tk.X, padx=20)
        
        self.label_arquivo = tk.Label(self.frame_arquivo, text="Nome do Arquivo:", 
                font=("Arial", 10))
        self.label_arquivo.pack(side=tk.LEFT)
        
        self.entry_arquivo = tk.Entry(self.frame_arquivo, width=30, font=("Arial", 10))
        self.entry_arquivo.pack(side=tk.LEFT, padx=10)
        self.entry_arquivo.insert(0, "caca_palavras")
        
        # Dropdown para formato
        self.label_formato = tk.Label(self.frame_arquivo, text="Formato:", 
                font=("Arial", 10))
        self.label_formato.pack(side=tk.LEFT, padx=(10, 5))
        
        # Verificar formatos disponíveis
        formatos_disponiveis = [".pdf"]
        if PILLOW_DISPONIVEL:
            formatos_disponiveis.append(".jpeg")
        if DOCX_DISPONIVEL:
            formatos_disponiveis.append(".docx")
        
        self.formato_var = tk.StringVar(value=".pdf")
        self.combo_formato = ttk.Combobox(self.frame_arquivo, textvariable=self.formato_var,
                                         values=formatos_disponiveis,
                                         state="readonly", width=10, font=("Arial", 10))
        self.combo_formato.pack(side=tk.LEFT)
        
        # Label informativo sobre bibliotecas faltantes
        if not PILLOW_DISPONIVEL or not DOCX_DISPONIVEL:
            msg_faltantes = []
            if not PILLOW_DISPONIVEL:
                msg_faltantes.append("JPEG (instale: pip install Pillow)")
            if not DOCX_DISPONIVEL:
                msg_faltantes.append("DOCX (instale: pip install python-docx)")
            
            self.label_info_libs = tk.Label(root, 
                text=f"⚠️ Formatos indisponíveis: {', '.join(msg_faltantes)}", 
                font=("Arial", 8), fg="orange", wraplength=550, justify=tk.LEFT)
            self.label_info_libs.pack(padx=20, pady=(0, 5))
        
        # Frame para seleção de diretório
        self.frame_destino = tk.Frame(root, pady=5)
        self.frame_destino.pack(fill=tk.X, padx=20)
        
        self.label_salvar = tk.Label(self.frame_destino, text="Salvar em:", 
                font=("Arial", 10))
        self.label_salvar.pack(side=tk.LEFT)
        
        self.label_destino = tk.Label(self.frame_destino, text=self.diretorio_destino, 
                                      font=("Arial", 9), 
                                      anchor=tk.W, width=35)
        self.label_destino.pack(side=tk.LEFT, padx=10)
        
        self.btn_escolher_pasta = tk.Button(self.frame_destino, text="Escolher Pasta", 
                                            font=("Arial", 9),
                                            command=self.escolher_pasta)
        self.btn_escolher_pasta.pack(side=tk.LEFT)
        
        # Botão gerar
        self.btn_gerar = tk.Button(root, text="GERAR CAÇA-PALAVRAS", 
                                   font=("Arial", 12, "bold"),
                                   bg="#4CAF50", fg="white",
                                   pady=10, cursor="hand2",
                                   command=self.gerar_caca_palavras)
        self.btn_gerar.pack(pady=15, padx=20, fill=tk.X)
        
        # Aplicar tema inicial
        self.aplicar_tema()
    
    def aplicar_tema(self):
        """Aplica o tema (claro ou escuro) em todos os widgets"""
        cores = self.cores_escuro if self.tema_escuro.get() else self.cores_claro
        
        # Janela principal
        self.root.config(bg=cores['bg'])
        
        # Frames do topo
        self.frame_topo.config(bg=cores['bg'])
        self.frame_titulo.config(bg=cores['bg'])
        
        # Títulos
        self.titulo.config(bg=cores['bg'], fg=cores['fg'])
        self.subtitulo.config(bg=cores['bg'], fg=cores['fg'])
        
        # Botão tema - configuração especial para não mostrar caixa branca
        self.btn_tema.config(
            bg=cores['bg'], 
            fg=cores['fg'],
            activebackground=cores['bg'],
            activeforeground=cores['fg'],
            selectcolor=cores['bg'],
            highlightbackground=cores['bg'],
            highlightcolor=cores['bg']
        )
        
        # Frames
        self.frame_tamanho.config(bg=cores['bg'])
        self.frame_opcoes.config(bg=cores['bg'], fg=cores['fg'], 
                                highlightbackground=cores['border'],
                                highlightcolor=cores['border'],
                                bd=2)
        self.frame_palavras.config(bg=cores['bg'])
        self.frame_arquivo.config(bg=cores['bg'])
        self.frame_destino.config(bg=cores['bg'])
        
        # Labels
        self.label_tamanho.config(bg=cores['bg'], fg=cores['fg'])
        self.label_exemplo_tamanho.config(bg=cores['bg'], fg=cores['label_info'])
        self.label_explicativo.config(bg=cores['bg'], fg=cores['label_info'])
        self.label_palavras.config(bg=cores['bg'], fg=cores['fg'])
        self.label_arquivo.config(bg=cores['bg'], fg=cores['fg'])
        self.label_formato.config(bg=cores['bg'], fg=cores['fg'])
        self.label_salvar.config(bg=cores['bg'], fg=cores['fg'])
        self.label_destino.config(bg=cores['bg'], fg=cores['label_link'])
        
        # Checkbuttons - configuração especial para não mostrar caixa branca
        self.check_diagonais.config(
            bg=cores['bg'], 
            fg=cores['fg'],
            activebackground=cores['bg'],
            activeforeground=cores['fg'],
            selectcolor=cores['bg'],
            highlightbackground=cores['bg']
        )
        self.check_contrarias.config(
            bg=cores['bg'], 
            fg=cores['fg'],
            activebackground=cores['bg'],
            activeforeground=cores['fg'],
            selectcolor=cores['bg'],
            highlightbackground=cores['bg']
        )
        
        # Entries
        self.entry_tamanho.config(bg=cores['entry_bg'], fg=cores['entry_fg'], 
                                  insertbackground=cores['entry_fg'],
                                  relief=tk.FLAT, bd=2,
                                  highlightthickness=1, highlightbackground=cores['border'],
                                  highlightcolor=cores['border'])
        self.entry_arquivo.config(bg=cores['entry_bg'], fg=cores['entry_fg'], 
                                 insertbackground=cores['entry_fg'],
                                 relief=tk.FLAT, bd=2,
                                 highlightthickness=1, highlightbackground=cores['border'],
                                 highlightcolor=cores['border'])
        
        # Text widget
        self.text_palavras.config(bg=cores['entry_bg'], fg=cores['entry_fg'], 
                                 insertbackground=cores['entry_fg'],
                                 relief=tk.FLAT, bd=2,
                                 highlightthickness=1, highlightbackground=cores['border'],
                                 highlightcolor=cores['border'])
        
        # Botões
        self.btn_escolher_pasta.config(bg=cores['button_bg'], fg=cores['button_fg'],
                                       activebackground=cores['button_bg'],
                                       relief=tk.FLAT, bd=0,
                                       highlightthickness=0)
        self.btn_gerar.config(bg=cores['button_bg'], fg=cores['button_fg'],
                             activebackground=cores['button_bg'],
                             relief=tk.FLAT, bd=0,
                             highlightthickness=0)
    
    def alternar_tema(self):
        """Alterna entre tema claro e escuro"""
        self.aplicar_tema()
    
    def escolher_pasta(self):
        """Abre o diálogo para escolher a pasta de destino"""
        pasta = filedialog.askdirectory(initialdir=self.diretorio_destino,
                                        title="Escolha onde salvar o PDF")
        if pasta:
            self.diretorio_destino = pasta
            # Truncar o caminho se for muito longo
            caminho_exibido = pasta
            if len(caminho_exibido) > 50:
                caminho_exibido = "..." + caminho_exibido[-47:]
            self.label_destino.config(text=caminho_exibido)
        
    def gerar_caca_palavras(self):
        """Função chamada ao clicar no botão gerar"""
        try:
            # Validar tamanho da grade
            tamanho_texto = self.entry_tamanho.get().strip()
            if 'x' in tamanho_texto.lower():
                partes = tamanho_texto.lower().split('x')
                tamanho = int(partes[0].strip())
            else:
                tamanho = int(tamanho_texto)
            
            if tamanho < 5 or tamanho > 30:
                messagebox.showerror("Erro", "O tamanho da grade deve estar entre 5 e 30")
                return
            
            # Obter palavras
            texto_palavras = self.text_palavras.get("1.0", tk.END)
            palavras = [p.strip() for p in texto_palavras.split('\n') if p.strip()]
            
            if len(palavras) == 0:
                messagebox.showerror("Erro", "Por favor, insira pelo menos uma palavra")
                return
            
            # Obter nome do arquivo
            nome_arquivo = self.entry_arquivo.get().strip()
            formato = self.formato_var.get()
            
            # Adicionar extensão se não tiver
            if not nome_arquivo.endswith(formato):
                if '.' in nome_arquivo:
                    nome_arquivo = nome_arquivo.rsplit('.', 1)[0]
                nome_arquivo += formato
            
            # Criar o caminho completo
            caminho_completo = os.path.join(self.diretorio_destino, nome_arquivo)
            
            # Obter opções
            usar_diagonais = self.var_diagonais.get()
            usar_contrarias = self.var_contrarias.get()
            
            # Gerar caça-palavras
            self.btn_gerar.config(state=tk.DISABLED, text="GERANDO...")
            self.root.update()
            
            gerador = GeradorCacaPalavras()
            gerador.criar_grade_vazia(tamanho)
            
            palavras_nao_inseridas = gerador.inserir_palavras(palavras, usar_diagonais, usar_contrarias)
            
            if palavras_nao_inseridas:
                resposta = messagebox.askyesno(
                    "Aviso",
                    f"As seguintes palavras não puderam ser inseridas:\n\n" +
                    "\n".join(palavras_nao_inseridas) +
                    "\n\nDeseja continuar mesmo assim?"
                )
                if not resposta:
                    self.btn_gerar.config(state=tk.NORMAL, text="GERAR CAÇA-PALAVRAS")
                    return
            
            gerador.preencher_espacos_vazios()
            
            # Gerar arquivo no formato escolhido
            if formato == '.pdf':
                gerador.gerar_pdf(caminho_completo, palavras)
            elif formato == '.jpeg':
                if not PILLOW_DISPONIVEL:
                    raise ImportError("Biblioteca Pillow não instalada. Use: pip install Pillow")
                gerador.gerar_jpeg(caminho_completo, palavras, incluir_gabarito=True)
                # Atualizar mensagem para JPEG (gera 2 arquivos)
                base_nome = caminho_completo.rsplit('.', 1)[0]
                caminho_completo = f"{base_nome}_caca.jpeg e {base_nome}_gabarito.jpeg"
            elif formato == '.docx':
                if not DOCX_DISPONIVEL:
                    raise ImportError("Biblioteca python-docx não instalada. Use: pip install python-docx")
                gerador.gerar_docx(caminho_completo, palavras)
            
            self.btn_gerar.config(state=tk.NORMAL, text="GERAR CAÇA-PALAVRAS")
            
            messagebox.showinfo("Sucesso", 
                              f"Caça-palavras gerado com sucesso!\n\n" +
                              f"Arquivo: {caminho_completo}\n" +
                              f"Palavras inseridas: {len(palavras) - len(palavras_nao_inseridas)}/{len(palavras)}")
            
        except ValueError as e:
            messagebox.showerror("Erro", "Por favor, insira um tamanho válido (ex: 18x18)")
            self.btn_gerar.config(state=tk.NORMAL, text="GERAR CAÇA-PALAVRAS")
        except ImportError as e:
            messagebox.showerror("Biblioteca Faltando", 
                               f"Erro: {str(e)}\n\n"
                               "Por favor, instale as bibliotecas necessárias usando:\n"
                               "pip install reportlab Pillow python-docx")
            self.btn_gerar.config(state=tk.NORMAL, text="GERAR CAÇA-PALAVRAS")
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao gerar o arquivo:\n\n{str(e)}")
            self.btn_gerar.config(state=tk.NORMAL, text="GERAR CAÇA-PALAVRAS")


if __name__ == "__main__":
    root = tk.Tk()
    app = InterfaceApp(root)
    root.mainloop()